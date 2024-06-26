import pandas as pd
from docx import Document
from docx.shared import Inches

# Load the spreadsheet-change it as it is
file_path = '/Users/yugoiwamoto/Desktop/IGM_project/IGM_incoming_list.xlsx'
try:
    data = pd.read_excel(file_path)
    print("Columns in the DataFrame:", data.columns)
    print("First few rows of the DataFrame:")
    print(data.head())
except FileNotFoundError:
    print(f"The file at {file_path} was not found. Please check the file path.")
    exit()
except Exception as e:
    print(f"An error occurred while loading the Excel file: {e}")
    exit()

# Trim whitespace from column names
data.columns = data.columns.str.strip()

# Create a new Word document
doc = Document()

# Avery 5160 layout: 3 columns, 10 rows per page
num_columns = 3
num_rows = 10

# Dimensions for Avery 5160 labels
label_width = 2.625  # in inches
label_height = 1.0   # in inches

# Add a table to the document
table = doc.add_table(rows=num_rows, cols=num_columns)
table.autofit = False

# Set the width of each column
for col in table.columns:
    for cell in col.cells:
        cell.width = Inches(label_width)

# Populate the table with names and placeholder for program
row_idx = 0
col_idx = 0

for index, row in data.iterrows():
    first_name = row.get('First Name', 'N/A')
    last_name = row.get('Last Name', 'N/A')
    program = 'Placeholder for Program'  # Placeholder text for program

    # Combine the names and program
    name_text = f"{first_name} {last_name}\nProgram: {program}"
    
    cell = table.cell(row_idx, col_idx)
    cell.text = f"{name_text}"
    
    # Add a new paragraph to adjust alignment
    cell.add_paragraph("\n")

    col_idx += 1
    if col_idx >= num_columns:
        col_idx = 0
        row_idx += 1
        if row_idx >= num_rows:
            # Add a page break and a new table if we exceed the row limit
            doc.add_page_break()
            table = doc.add_table(rows=num_rows, cols=num_columns)
            table.autofit = False
            for col in table.columns:
                for cell in col.cells:
                    cell.width = Inches(label_width)
            row_idx = 0

# Save the document
output_path = '/Users/yugoiwamoto/Desktop/IGM_project/IGM_incoming_name_tags.docx'
try:
    doc.save(output_path)
    print(f"Document has been created successfully at {output_path}.")
except Exception as e:
    print(f"An error occurred while saving the document: {e}")
